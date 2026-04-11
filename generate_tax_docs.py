"""세금 신고 관련 DOCX 3종 생성"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "docx_contracts"
os.makedirs(OUT, exist_ok=True)

def setup(doc):
    s = doc.styles['Normal']
    s.font.name = 'Malgun Gothic'
    s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.5

def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    doc.add_paragraph()

def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(12)

def legal(doc, refs):
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("[법적 근거]")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    for ref in refs:
        p2 = doc.add_paragraph(f"  {ref}")
        for run in p2.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88,0x88,0x88)

# ===== 7. 종합소득세 신고 준비 체크리스트 =====
def doc_tax_checklist():
    d = Document(); setup(d)
    title(d, "종합소득세 신고 준비 체크리스트")
    d.add_paragraph("신고 기간: 매년 5월 1일 ~ 5월 31일")
    d.add_paragraph("대상: 동호회 대표자 (사업소득)")
    d.add_paragraph()

    heading(d, "1단계: 수입 자료 정리 (4월 중 완료)")
    items = [
        "[ ] 단체 통장 연간 거래내역서 발급 (은행 방문 또는 인터넷뱅킹)",
        "[ ] 회비 수입 총액 계산 (월별 합계)",
        "[ ] 교회 월세 수입 총액 계산",
        "[ ] 관리 수수료 수입 총액 계산",
        "[ ] 기타 수입 (행사 참가비, 후원금 등) 총액",
        "[ ] 연간 총 수입 합계: ________________원",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "2단계: 비용(필요경비) 자료 정리 (4월 중 완료)")
    items = [
        "[ ] 카드 사용 내역서 발급 (연간) — 카드사 홈페이지",
        "[ ] 현금영수증 사용 내역 — 홈택스 조회",
        "[ ] 월별 영수증 봉투 12개 확인 (누락 없는지)",
        "[ ] 엑셀 장부 월간 합계 12개월 확인",
        "",
        "[ ] 장소 임대료 합계: ________________원 (증빙: 임대차 계약서 + 이체내역)",
        "[ ] 강사비 합계: ________________원 (증빙: 용역 계약서 + 이체내역)",
        "[ ] 원천징수세 납부 합계: ________________원 (증빙: 홈택스 납부 내역)",
        "[ ] 교통비 합계: ________________원 (증빙: 카드전표/하이패스 내역)",
        "[ ] 식비/회의비 합계: ________________원 (증빙: 카드전표)",
        "[ ] 사무용품/소모품 합계: ________________원 (증빙: 카드전표)",
        "[ ] 통신비 합계: ________________원 (증빙: 고지서/자동이체 내역)",
        "[ ] 교재/교육 자료 합계: ________________원 (증빙: 카드전표/영수증)",
        "[ ] 행사비 합계: ________________원 (증빙: 카드전표/영수증)",
        "[ ] 보험료 합계: ________________원 (증빙: 보험 납입 영수증)",
        "[ ] 수리비/유지보수 합계: ________________원 (증빙: 세금계산서/영수증)",
        "[ ] 기타 비용 합계: ________________원",
        "",
        "[ ] 연간 총 비용(필요경비) 합계: ________________원",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "3단계: 과세 소득 계산")
    items = [
        "총 수입:          ________________원 (A)",
        "총 필요경비:      ________________원 (B)",
        "과세표준:         ________________원 (A - B)",
        "",
        "적용 세율: (아래 표 참고)",
        "  1,400만 이하: 6%",
        "  1,400~5,000만: 15% (누진공제 126만)",
        "  5,000~8,800만: 24% (누진공제 576만)",
        "",
        "산출 세액:        ________________원",
        "지방소득세 (10%): ________________원",
        "총 납부 세액:     ________________원",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "4단계: 홈택스 신고 (5월 1~31일)")
    items = [
        "[ ] 홈택스 (www.hometax.go.kr) 접속",
        "[ ] 종합소득세 신고 → 일반신고 (또는 단순경비율 신고)",
        "[ ] 사업장 정보 입력 (고유번호증 번호)",
        "[ ] 수입금액 입력",
        "[ ] 필요경비 입력 (항목별)",
        "[ ] 세액 자동 계산 확인",
        "[ ] 신고서 제출",
        "[ ] 납부서 출력 → 은행 납부 또는 계좌이체",
        "",
        "[ ] 6월: 지방소득세 별도 납부 (위택스 또는 구청)",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "5단계: 서류 보관 (신고 후)")
    items = [
        "[ ] 신고서 사본 PDF 저장",
        "[ ] 납부 영수증 보관",
        "[ ] 엑셀 장부 + 영수증 봉투 5년간 보관",
        "[ ] 통장 거래내역서 5년간 보관",
        "",
        "보관 기간: 5년 (국세기본법 제26조의2)",
    ]
    for item in items:
        d.add_paragraph(item)

    legal(d, [
        "소득세법 제70조 (종합소득 과세표준 확정신고)",
        "소득세법 제55조 (세율)",
        "소득세법 제27조 (사업소득의 필요경비)",
        "국세기본법 제26조의2 (국세부과 제척기간 — 5년)",
        "지방세법 제89조 (개인지방소득세)",
    ])
    d.save(f"{OUT}/07_Tax_Filing_Checklist.docx")
    print("7/9 종합소득세 체크리스트")


# ===== 8. 원천징수 이행상황 신고 가이드 =====
def doc_withholding():
    d = Document(); setup(d)
    title(d, "원천징수 이행상황 신고 가이드")
    d.add_paragraph("신고 기한: 매월 10일 (반기납 선택 시 1/10, 7/10)")
    d.add_paragraph("대상: 강사비 등 사업소득 지급 시 3.3% 원천징수")
    d.add_paragraph()

    heading(d, "1. 원천징수 계산 방법")
    items = [
        "강사비 지급액:     ________________원 (A)",
        "소득세 (3%):       ________________원 (A x 0.03)",
        "지방소득세 (0.3%): ________________원 (A x 0.003)",
        "원천징수 합계:     ________________원 (A x 0.033)",
        "강사 실수령액:     ________________원 (A - 원천징수 합계)",
        "",
        "예시: 강사비 80만 원",
        "  소득세: 800,000 x 3% = 24,000원",
        "  지방소득세: 800,000 x 0.3% = 2,400원",
        "  원천징수 합계: 26,400원",
        "  강사 실수령액: 773,600원",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "2. 신고 방법 (홈택스)")
    items = [
        "[ ] 홈택스 접속 → 신고/납부 → 원천세",
        "[ ] 원천징수이행상황신고서 작성",
        "[ ] 사업소득(기타소득) 란에 지급총액, 소득세, 지방소득세 입력",
        "[ ] 전자납부 또는 가상계좌 납부",
        "",
        "반기납 신청 방법:",
        "[ ] 홈택스 → 신청/제출 → 원천징수세액 반기별 납부 승인신청",
        "[ ] 승인 후 1~6월분은 7/10까지, 7~12월분은 다음해 1/10까지 신고",
    ]
    for item in items:
        d.add_paragraph(item)

    heading(d, "3. 월별 원천징수 기록표")
    table = d.add_table(rows=14, cols=6)
    headers = ["월", "강사명", "지급액", "소득세(3%)", "지방세(0.3%)", "실수령액"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        for run in table.cell(0, j).paragraphs[0].runs:
            run.bold = True
    for i in range(1, 13):
        table.cell(i, 0).text = f"{i}월"
    table.cell(13, 0).text = "합계"
    for run in table.cell(13, 0).paragraphs[0].runs:
        run.bold = True

    heading(d, "4. 지급명세서 제출")
    items = [
        "[ ] 매년 3월 10일까지 전년도 지급명세서 제출 (홈택스)",
        "[ ] 사업소득 지급명세서 (강사별 연간 지급 총액)",
        "[ ] 미제출 시 가산세: 지급금액의 2%",
    ]
    for item in items:
        d.add_paragraph(item)

    legal(d, [
        "소득세법 제127조 (원천징수의무)",
        "소득세법 제129조 (원천징수세율)",
        "소득세법 제164조 (지급명세서 제출)",
        "지방세법 제95조의2 (특별징수)",
    ])
    d.save(f"{OUT}/08_Withholding_Tax_Guide.docx")
    print("8/9 원천징수 가이드")


# ===== 9. 연간 세금 캘린더 + 준비 서류 =====
def doc_tax_calendar():
    d = Document(); setup(d)
    title(d, "연간 세금 신고 캘린더 및 준비 서류 목록")
    d.add_paragraph("비영리단체(동호회) 대표자용 — 2026년 기준")
    d.add_paragraph()

    heading(d, "월별 세금 일정")
    table = d.add_table(rows=13, cols=4)
    headers = ["월", "신고/납부 항목", "마감일", "준비 서류"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        for run in table.cell(0, j).paragraphs[0].runs:
            run.bold = True

    data = [
        ("1월", "원천세 반기납 (7~12월분)\n정기총회 개최", "1/10\n1월 중", "원천징수이행상황신고서\n전년도 결산보고서"),
        ("2월", "전년도 결산 완료\n감사 보고", "2월 말", "수입지출 결산서\n감사보고서"),
        ("3월", "지급명세서 제출\n총회 보고 마감", "3/10", "사업소득 지급명세서\n총회 회의록"),
        ("4월", "종합소득세 준비\n자료 정리", "4월 말", "통장내역서, 카드내역서\n영수증 정리 완료"),
        ("5월", "★ 종합소득세 신고 ★", "5/31", "종합소득세 신고서\n필요경비 증빙 일체"),
        ("6월", "지방소득세 납부", "6/30", "지방소득세 납부서"),
        ("7월", "원천세 반기납 (1~6월분)", "7/10", "원천징수이행상황신고서"),
        ("8월", "하반기 예산 점검", "—", "상반기 정산표"),
        ("9월", "—", "—", "—"),
        ("10월", "연말 비용 소진 계획", "—", "비용 사용률 점검표"),
        ("11월", "비용 집중 집행\n물품 구매/수리", "—", "견적서, 영수증"),
        ("12월", "연간 결산 준비\n회비 완납 확인", "12월 말", "연간 수입지출표\n회비 완납 확인서"),
    ]
    for i, (month, item, deadline, docs) in enumerate(data):
        table.cell(i+1, 0).text = month
        table.cell(i+1, 1).text = item
        table.cell(i+1, 2).text = deadline
        table.cell(i+1, 3).text = docs

    heading(d, "준비 서류 종합 체크리스트 (13단계)")
    table2 = d.add_table(rows=14, cols=5)
    headers2 = ["#", "서류명", "발급처", "비용", "소요기간"]
    for j, h in enumerate(headers2):
        table2.cell(0, j).text = h
        for run in table2.cell(0, j).paragraphs[0].runs:
            run.bold = True

    docs_data = [
        ("1", "고유번호증", "관할 세무서", "무료", "1~2주"),
        ("2", "정관(규약)", "자체 작성", "무료", "1일"),
        ("3", "설립총회 회의록", "자체 작성", "무료", "1시간"),
        ("4", "단체 통장", "은행", "무료", "1시간"),
        ("5", "단체 체크카드", "은행", "무료", "1주"),
        ("6", "임대차 계약서", "자체 작성+서명", "무료", "30분"),
        ("7", "관리위탁 계약서", "자체 작성+서명", "무료", "30분"),
        ("8", "전대차 계약서", "자체 작성+서명", "무료", "30분"),
        ("9", "교육용역 계약서", "자체 작성+서명", "무료", "20분"),
        ("10", "사업자용 현금영수증 등록", "홈택스", "무료", "10분"),
        ("11", "원천징수 반기납 신청", "홈택스", "무료", "10분"),
        ("12", "엑셀/구글시트 장부", "자체 생성", "무료", "30분"),
        ("13", "영수증 보관함 (봉투 12개)", "문구점", "약 5,000원", "10분"),
    ]
    for i, row in enumerate(docs_data):
        for j, val in enumerate(row):
            table2.cell(i+1, j).text = val

    d.add_paragraph()
    d.add_paragraph("총 비용: 약 5,000원 (영수증 봉투)")
    d.add_paragraph("총 소요기간: 약 2~3주 (세무서 처리 기간 포함)")

    legal(d, [
        "소득세법 제70조 (종합소득세 신고)",
        "소득세법 제127조 (원천징수)",
        "국세기본법 시행규칙 제7조 (고유번호 부여)",
        "국세기본법 제26조의2 (서류 보관 5년)",
    ])
    d.save(f"{OUT}/09_Tax_Calendar_Documents.docx")
    print("9/9 세금 캘린더+준비서류")


if __name__ == "__main__":
    doc_tax_checklist()
    doc_withholding()
    doc_tax_calendar()
    print("모든 세금 DOCX 생성 완료!")
