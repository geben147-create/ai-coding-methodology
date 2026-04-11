"""백송한신 저가매매 관련 DOCX 파일 생성 스크립트"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run.bold = bold


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
    return h


def add_para(doc, text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


# =====================================================
# 문서 1: 부동산 매매계약서
# =====================================================
def create_contract():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    add_para(doc, "부동산 매매계약서", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "", size=8)

    # 당사자
    add_heading(doc, "당사자", level=2)
    add_para(doc, "매도인(이하 \"갑\"): 홍○○ (주민등록번호: ______-______)")
    add_para(doc, "주소: ")
    add_para(doc, "")
    add_para(doc, "매수인(이하 \"을\"): 홍○○ (주민등록번호: ______-______)")
    add_para(doc, "주소: ")
    add_para(doc, "")
    add_para(doc, "갑과 을은 부녀 관계(직계존비속)로서, 아래 부동산에 관하여 다음과 같이 매매계약을 체결한다.")
    add_para(doc, "")

    # 제1조
    add_heading(doc, "제1조 (매매 목적물)", level=2)
    table1 = doc.add_table(rows=6, cols=2, style="Table Grid")
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ("소재지", "경기도 고양시 일산동구 백석동 ○○○"),
        ("단지명", "백송한신아파트 ○○○동 ○○○호"),
        ("면적", "전용 37.○○㎡ (공급면적 50㎡ B타입)"),
        ("구조", "철근콘크리트 / 아파트"),
        ("이전 범위", "☐ 전체 지분  ☐ 50% 지분 (해당란에 체크)"),
        ("현 임대차", "전세 / 보증금: ______만원 / 만기: 20__년 __월 __일"),
    ]
    for i, (k, v) in enumerate(items):
        set_cell_text(table1.cell(i, 0), k, bold=True)
        set_cell_text(table1.cell(i, 1), v)
    add_para(doc, "")

    # 제2조
    add_heading(doc, "제2조 (매매대금)", level=2)
    add_para(doc, "매매대금 총액: 금 일억오천만원정 (₩150,000,000)", bold=True)
    add_para(doc, "")

    table2 = doc.add_table(rows=4, cols=3, style="Table Grid")
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["구분", "금액", "지급일"]
    for i, h in enumerate(headers):
        set_cell_text(table2.cell(0, i), h, bold=True)
    rows = [
        ("계약금", "금 ______만원정", "계약 체결일 (20__년 __월 __일)"),
        ("중도금", "금 ______만원정 (해당 시)", "20__년 __월 __일"),
        ("잔  금", "금 ______만원정", "20__년 __월 __일"),
    ]
    for i, (a, b, c) in enumerate(rows):
        set_cell_text(table2.cell(i+1, 0), a)
        set_cell_text(table2.cell(i+1, 1), b)
        set_cell_text(table2.cell(i+1, 2), c)
    add_para(doc, "")
    add_para(doc, "단, 잔금 중 금 ______만원은 제3조의 기존 채무 상계로 갈음한다.")
    add_para(doc, "모든 대금 지급은 반드시 계좌이체로 하며, 현금 지급은 인정하지 않는다.")
    add_para(doc, "")

    # 제3조
    add_heading(doc, "제3조 (기존 채무 상계)", level=2)
    add_para(doc, "① 을이 갑에게 보유하고 있는 기존 채권 금 ______만원(별첨 「채무확인 및 상계합의서」 참조)을 본 매매대금의 일부와 상계한다.")
    add_para(doc, "② 상계 후 실제 지급할 잔금: 금 ______만원")
    add_para(doc, "③ 상계 처리되는 채권의 근거는 별첨 은행 계좌이체 내역 및 채무확인서로 한다.")
    add_para(doc, "")

    # 제4조
    add_heading(doc, "제4조 (매매가액 산정 근거)", level=2)
    add_para(doc, "본 매매가액은 다음 사항을 종합적으로 고려하여 산정하였다:")
    add_para(doc, "  1) 국토교통부 실거래가 공개시스템 기준: 2026년 1월 동일 단지 동일 평형 실거래가 2억 3,000만원 (12층)")
    add_para(doc, "  2) 매매 목적물의 층수, 향, 내부 상태에 따른 감가 요인")
    add_para(doc, "  3) 갑·을 간 기존 채권채무 관계 반영")
    add_para(doc, "  4) 상속세 및 증여세법 제35조(저가 양수에 따른 이익의 증여) 기준 검토")
    add_para(doc, "  5) 시가 대비 차액이 시가의 30% 이내 또는 3억원 이내 기준 준수 여부 확인")
    add_para(doc, "")

    # 제5조
    add_heading(doc, "제5조 (소유권 이전)", level=2)
    add_para(doc, "① 갑은 잔금 수령과 동시에 을 명의로 소유권 이전 등기에 필요한 일체의 서류를 교부한다.")
    add_para(doc, "② 소유권 이전 등기에 소요되는 비용(취득세, 등록세, 법무사 수수료 등)은 을이 부담한다.")
    add_para(doc, "③ 갑은 소유권 이전 시까지 목적물에 새로운 제한물권을 설정하거나 부담을 지우지 않는다.")
    add_para(doc, "")

    # 제6조
    add_heading(doc, "제6조 (임대차 승계)", level=2)
    add_para(doc, "① 목적물에 현존하는 임대차(전세)계약은 을이 임대인 지위를 승계한다.")
    add_para(doc, "② 임차인에 대한 보증금 반환 의무는 소유권 이전일부터 을이 부담한다.")
    add_para(doc, "")

    # 제7조
    add_heading(doc, "제7조 (하자 면책)", level=2)
    add_para(doc, "을은 목적물의 현 상태를 충분히 확인하였으며, 물리적 하자에 대하여 갑에게 책임을 묻지 않는다.")
    add_para(doc, "")

    # 제8조
    add_heading(doc, "제8조 (특약사항)", level=2)
    add_para(doc, "  1) 본 거래는 직계존비속 간 매매로서, 상속세 및 증여세법 관련 사항을 충분히 인지한 상태에서 체결한다.")
    add_para(doc, "  2) ")
    add_para(doc, "  3) ")
    add_para(doc, "")

    # 서명란
    add_para(doc, "")
    add_para(doc, "위 계약의 성립을 증명하기 위하여 본 계약서 2통을 작성하고, 갑·을이 서명 날인한 후 각 1통씩 보관한다.",
             size=10)
    add_para(doc, "")
    add_para(doc, "2026년      월      일", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "")
    add_para(doc, "매도인(갑):                              (인)", size=12)
    add_para(doc, "매수인(을):                              (인)", size=12)

    path = os.path.join(DOCS_DIR, "01_부동산매매계약서_백송한신.docx")
    doc.save(path)
    print(f"Created: {path}")


# =====================================================
# 문서 2: 채무확인 및 상계합의서
# =====================================================
def create_debt_agreement():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    add_para(doc, "채무확인 및 상계합의서", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "", size=8)

    # 당사자
    add_heading(doc, "당사자", level=2)
    add_para(doc, "채무자(이하 \"갑\"): 홍○○ (아버지) (주민등록번호: ______-______)")
    add_para(doc, "채권자(이하 \"을\"): 홍○○ (딸)   (주민등록번호: ______-______)")
    add_para(doc, "")
    add_para(doc, "갑과 을은 아래와 같이 기존 채무를 확인하고, 부동산 매매대금과의 상계에 합의한다.")
    add_para(doc, "")

    # 제1조
    add_heading(doc, "제1조 (갑의 을에 대한 채무 확인)", level=2)
    add_para(doc, "갑은 을에게 아래 사유로 총 금 ______만원의 채무가 있음을 확인한다:")
    add_para(doc, "")

    table1 = doc.add_table(rows=5, cols=5, style="Table Grid")
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["번호", "일자(기간)", "사유", "금액(만원)", "증빙"]):
        set_cell_text(table1.cell(0, i), h, bold=True, size=9)
    sample = [
        ("1", "2020.○○~2024.○○", "을→갑 계좌이체 (생활지원/대여)", "", "별첨 1: 은행내역"),
        ("2", "2021.○○", "을→갑 암호화폐 투자위탁 원금", "", "별첨 2: 거래소내역"),
        ("3", "", "", "", ""),
        ("합계", "", "", "", ""),
    ]
    for i, row in enumerate(sample):
        for j, val in enumerate(row):
            set_cell_text(table1.cell(i+1, j), val, size=9, bold=(i==3))
    add_para(doc, "")

    # 제2조
    add_heading(doc, "제2조 (갑의 을에 대한 반대채권)", level=2)
    add_para(doc, "을은 갑으로부터 아래 금원을 수령하였음을 확인한다:")
    add_para(doc, "")

    table2 = doc.add_table(rows=5, cols=5, style="Table Grid")
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["번호", "일자(기간)", "사유", "금액(만원)", "증빙"]):
        set_cell_text(table2.cell(0, i), h, bold=True, size=9)
    sample2 = [
        ("1", "20__년 ○월", "갑→을 창업지원/증여", "", "은행내역"),
        ("2", "20__년 ○월", "투자수익 분배/환원", "", "거래소내역"),
        ("3", "", "", "", ""),
        ("합계", "", "", "", ""),
    ]
    for i, row in enumerate(sample2):
        for j, val in enumerate(row):
            set_cell_text(table2.cell(i+1, j), val, size=9, bold=(i==3))
    add_para(doc, "")

    # 제3조
    add_heading(doc, "제3조 (순채무액 확정)", level=2)
    add_para(doc, "제1조의 채무 합계에서 제2조의 반대채권을 차감한 갑의 순채무:")
    add_para(doc, "")
    add_para(doc, "  순채무 = 제1조 합계 ______만원 - 제2조 합계 ______만원 = 금 ______만원", bold=True, size=12)
    add_para(doc, "")

    # 제4조
    add_heading(doc, "제4조 (상계 합의)", level=2)
    add_para(doc, "① 갑과 을은 위 순채무 금 ______만원을 별도 체결하는 「부동산 매매계약서」(백송한신아파트 ○○○동 ○○○호)의 매매대금 일부와 상계하기로 합의한다.")
    add_para(doc, "② 상계 후 을이 갑에게 실제 지급할 매매 잔금은 금 ______만원으로 한다.")
    add_para(doc, "③ 상계 처리 후 갑의 을에 대한 잔존 채무는 없는 것으로 한다. (또는: 잔존 채무 ______만원은 향후 월 ______만원씩 분할 상환한다.)")
    add_para(doc, "")

    # 제5조
    add_heading(doc, "제5조 (증빙서류)", level=2)
    add_para(doc, "본 합의서에 첨부하는 증빙서류는 다음과 같다:")
    add_para(doc, "  별첨 1: 은행 계좌이체 내역서 (○○은행, ○○은행) — 을→갑 이체 내역")
    add_para(doc, "  별첨 2: 바이낸스(Binance) 거래소 입출금 내역 및 손익 보고서")
    add_para(doc, "  별첨 3: 기존 차용증 사본 (해당 시)")
    add_para(doc, "  별첨 4: 은행 계좌이체 내역서 — 갑→을 이체 내역")
    add_para(doc, "")

    # 제6조
    add_heading(doc, "제6조 (확인 및 보증)", level=2)
    add_para(doc, "① 갑과 을은 본 합의서에 기재된 채무 내역이 사실과 부합함을 상호 확인한다.")
    add_para(doc, "② 본 합의서에 기재되지 않은 별도의 채권채무는 존재하지 않음을 상호 보증한다.")
    add_para(doc, "")

    # 서명란
    add_para(doc, "")
    add_para(doc, "위 사항을 확인하고 합의하기 위하여 본 합의서 2통을 작성하고, 갑·을이 서명 날인한 후 각 1통씩 보관한다.", size=10)
    add_para(doc, "")
    add_para(doc, "2026년      월      일", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "")
    add_para(doc, "채무자(갑):                              (인)", size=12)
    add_para(doc, "채권자(을):                              (인)", size=12)

    path = os.path.join(DOCS_DIR, "02_채무확인_상계합의서.docx")
    doc.save(path)
    print(f"Created: {path}")


# =====================================================
# 문서 3: 자금흐름 분석 정리표
# =====================================================
def create_fund_flow():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    add_para(doc, "자금 흐름 분석 정리표", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "백송한신아파트 저가매매를 위한 부녀간 자금 흐름 증빙 자료", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")

    # 섹션 1: 딸→아버지
    add_heading(doc, "1. 딸 → 아버지 자금 이전 내역", level=2)
    add_para(doc, "아래 표에 은행 계좌이체 내역을 기준으로 기입합니다. (계좌이체 확인서 별첨)")
    add_para(doc, "")

    t1 = doc.add_table(rows=8, cols=6, style="Table Grid")
    for i, h in enumerate(["No.", "이체일자", "출금계좌(딸)", "입금계좌(아버지)", "금액(원)", "성격 (대여/증여/투자/기타)"]):
        set_cell_text(t1.cell(0, i), h, bold=True, size=9)
    for r in range(1, 7):
        set_cell_text(t1.cell(r, 0), str(r), size=9)
    set_cell_text(t1.cell(7, 0), "합계", bold=True, size=9)
    set_cell_text(t1.cell(7, 4), "₩", bold=True, size=9)
    add_para(doc, "")

    # 섹션 2: 아버지→딸
    add_heading(doc, "2. 아버지 → 딸 자금 이전 내역", level=2)
    t2 = doc.add_table(rows=8, cols=6, style="Table Grid")
    for i, h in enumerate(["No.", "이체일자", "출금계좌(아버지)", "입금계좌(딸)", "금액(원)", "성격 (증여/투자환원/생활비/기타)"]):
        set_cell_text(t2.cell(0, i), h, bold=True, size=9)
    for r in range(1, 7):
        set_cell_text(t2.cell(r, 0), str(r), size=9)
    set_cell_text(t2.cell(7, 0), "합계", bold=True, size=9)
    set_cell_text(t2.cell(7, 4), "₩", bold=True, size=9)
    add_para(doc, "")

    # 섹션 3: 순이전액 계산
    add_heading(doc, "3. 순이전액 계산", level=2)
    t3 = doc.add_table(rows=4, cols=2, style="Table Grid")
    calc = [
        ("① 딸→아버지 합계", "₩"),
        ("② 아버지→딸 합계", "₩"),
        ("③ 순이전액 (① - ②)", "₩"),
        ("④ 상계 적용 가능 금액", "₩"),
    ]
    for i, (k, v) in enumerate(calc):
        set_cell_text(t3.cell(i, 0), k, bold=True)
        set_cell_text(t3.cell(i, 1), v)
    add_para(doc, "")

    # 섹션 4: 바이낸스
    add_heading(doc, "4. 암호화폐(바이낸스) 공동 운용 내역", level=2)
    t4 = doc.add_table(rows=6, cols=5, style="Table Grid")
    for i, h in enumerate(["구분", "일자", "금액(원환산)", "비고", "증빙"]):
        set_cell_text(t4.cell(0, i), h, bold=True, size=9)
    crypto_rows = [
        ("딸 출자 원금", "", "₩", "", "거래소 입금내역"),
        ("아버지 출자 원금", "", "₩", "", "거래소 입금내역"),
        ("현재 잔고 평가액", "2026.04 기준", "₩", "", "거래소 스크린샷"),
        ("실현 손익", "", "₩", "수익(+) / 손실(-)", "거래소 PnL 리포트"),
        ("딸 몫 정산액", "", "₩", "출자비율 적용", ""),
    ]
    for i, row in enumerate(crypto_rows):
        for j, val in enumerate(row):
            set_cell_text(t4.cell(i+1, j), val, size=9)
    add_para(doc, "")

    # 섹션 5: 요약
    add_heading(doc, "5. 최종 요약", level=2)
    t5 = doc.add_table(rows=5, cols=2, style="Table Grid")
    summary = [
        ("A. 은행 기준 순이전액 (딸→아버지)", "₩"),
        ("B. 암호화폐 딸 몫 미정산액", "₩"),
        ("C. 총 상계 가능 금액 (A + B)", "₩"),
        ("D. 매매대금", "₩150,000,000"),
        ("E. 실제 지급 필요액 (D - C)", "₩"),
    ]
    for i, (k, v) in enumerate(summary):
        set_cell_text(t5.cell(i, 0), k, bold=True)
        set_cell_text(t5.cell(i, 1), v)
    add_para(doc, "")

    add_para(doc, "작성일:  2026년    월    일", size=10)
    add_para(doc, "작성자:                        (서명)", size=10)

    path = os.path.join(DOCS_DIR, "03_자금흐름_분석정리표.docx")
    doc.save(path)
    print(f"Created: {path}")


# =====================================================
# 문서 4: 세금 시뮬레이션 요약
# =====================================================
def create_tax_summary():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    add_para(doc, "저가매매 세금 시뮬레이션 요약", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    add_para(doc, "백송한신아파트 50B㎡(전용37B) | 2026.04 기준", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")

    # 기본 정보
    add_heading(doc, "1. 기본 정보", level=2)
    t0 = doc.add_table(rows=6, cols=2, style="Table Grid")
    basics = [
        ("대상 물건", "백송한신아파트 ○○동 ○○호 (50B㎡, 전용37B)"),
        ("매도인", "아버지 (홍○○)"),
        ("매수인", "딸 (홍○○)"),
        ("관계", "직계존비속 (특수관계인)"),
        ("시가 (국토부 실거래)", "2억 3,000만원 (2026.01 실거래 기준)"),
        ("희망 매수가", "1억 5,000만원"),
    ]
    for i, (k, v) in enumerate(basics):
        set_cell_text(t0.cell(i, 0), k, bold=True)
        set_cell_text(t0.cell(i, 1), v)
    add_para(doc, "")

    # 법적 근거
    add_heading(doc, "2. 법적 근거", level=2)
    add_para(doc, "• 상속세 및 증여세법 제35조: 저가 양수에 따른 이익의 증여")
    add_para(doc, "• 동법 시행령 제26조: 시가와 대가의 차액 > MIN(시가×30%, 3억원) → 과세")
    add_para(doc, "• 동법 제53조: 직계존비속 증여재산 공제 5,000만원 (10년 합산)")
    add_para(doc, "• 동법 시행령 제49조: 시가 = 평가기준일 전후 6개월 이내 실거래가")
    add_para(doc, "")

    # 시뮬레이션
    add_heading(doc, "3. 매수가별 세금 시뮬레이션", level=2)
    add_para(doc, "시가 기준: 2억 3,000만원")
    add_para(doc, "")

    t1 = doc.add_table(rows=7, cols=7, style="Table Grid")
    headers = ["매수가", "차액", "시가30%", "과세여부", "증여재산가액", "증여세", "취득세(1.1%)"]
    for i, h in enumerate(headers):
        set_cell_text(t1.cell(0, i), h, bold=True, size=9)
    sim_data = [
        ("1억 5,000만", "8,000만", "6,900만", "과세 대상", "1,100만", "0원*", "165만"),
        ("1억 6,000만", "7,000만", "6,900만", "과세(근소)", "100만", "0원*", "176만"),
        ("1억 6,100만", "6,900만", "6,900만", "비과세", "0원", "0원", "177만"),
        ("1억 7,000만", "6,000만", "6,900만", "비과세", "0원", "0원", "187만"),
        ("1억 8,000만", "5,000만", "6,900만", "비과세", "0원", "0원", "198만"),
        ("2억", "3,000만", "6,900만", "비과세", "0원", "0원", "220만"),
    ]
    for i, row in enumerate(sim_data):
        for j, val in enumerate(row):
            bold = (i == 2)  # highlight safe line
            set_cell_text(t1.cell(i+1, j), val, size=9, bold=bold)

    add_para(doc, "")
    add_para(doc, "* 증여재산가액이 직계존비속 공제(5,000만원) 이내이므로 증여세 0원. 단, 10년 내 기존 증여 합산액 주의.", size=9)
    add_para(doc, "")

    # 권고안
    add_heading(doc, "4. 권고안", level=2)
    add_para(doc, "✅ 안전 최저가: 1억 6,100만원 (시가의 70%) — 증여세 비과세")
    add_para(doc, "⚠️ 1억 5,000만원 매수 시: 증여세 자체는 0원(공제 내)이나 과세 '대상'이 되어 세무 리스크 존재")
    add_para(doc, "❌ 증빙 없는 저가 매수: 세무조사 시 소명 곤란, 가산세 위험")
    add_para(doc, "")
    add_para(doc, "출처: 국토교통부 실거래가 공개시스템, 상속세 및 증여세법", size=9)
    add_para(doc, "주의: 이 문서는 세무/법률 자문이 아닌 분석 참고 자료입니다.", size=9, bold=True)

    path = os.path.join(DOCS_DIR, "04_세금시뮬레이션_요약.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_contract()
    create_debt_agreement()
    create_fund_flow()
    create_tax_summary()
    print("\nAll DOCX files created successfully!")
