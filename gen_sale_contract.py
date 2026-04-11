# -*- coding: utf-8 -*-
"""Generate DOCX: Contract + Offset Agreement for family low-price sale."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.dirname(os.path.abspath(__file__))

def sc(cell, text, bold=False, sz=10):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(sz)
            r.bold = bold

def heading(doc, text, sz=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sz)

def add_items(doc, lst, indent=1, sz=10):
    for i in lst:
        p = doc.add_paragraph(i)
        p.paragraph_format.left_indent = Cm(indent)
        for r in p.runs:
            r.font.size = Pt(sz)

def party_table(doc, role):
    tb = doc.add_table(rows=4, cols=4)
    tb.style = "Table Grid"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc(tb.rows[0].cells[0], role, bold=True)
    sc(tb.rows[0].cells[1], "\uc131 \uba85")
    sc(tb.rows[0].cells[2], "")
    sc(tb.rows[0].cells[3], "(\uc778)")
    sc(tb.rows[1].cells[0], "")
    sc(tb.rows[1].cells[1], "\uc8fc\ubbfc\ub4f1\ub85d\ubc88\ud638")
    sc(tb.rows[1].cells[2], "         -         ")
    sc(tb.rows[1].cells[3], "")
    sc(tb.rows[2].cells[0], "")
    sc(tb.rows[2].cells[1], "\uc8fc \uc18c")
    tb.rows[2].cells[2].merge(tb.rows[2].cells[3])
    sc(tb.rows[2].cells[2], "")
    sc(tb.rows[3].cells[0], "")
    sc(tb.rows[3].cells[1], "\uc5f0\ub77d\ucc98")
    sc(tb.rows[3].cells[2], "")
    sc(tb.rows[3].cells[3], "")

def sign_block(doc, roles):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("20      \ub144      \uc6d4      \uc77c")
    r.font.size = Pt(13)
    doc.add_paragraph("")
    doc.add_paragraph("")
    for role in roles:
        p = doc.add_paragraph()
        r = p.add_run(role)
        r.bold = True
        r.font.size = Pt(12)
        p = doc.add_paragraph()
        r = p.add_run("   \uc131 \uba85:                                        (\uc778)")
        r.font.size = Pt(12)
        p = doc.add_paragraph()
        r = p.add_run("   \uc8fc \uc18c:                                                              ")
        r.font.size = Pt(11)
        doc.add_paragraph("")

def init_doc():
    d = Document()
    s = d.styles["Normal"]
    s.font.size = Pt(11)
    s.paragraph_format.line_spacing = 1.8
    return d

def make_transfer_table(doc, rows=8):
    tt = doc.add_table(rows=rows+2, cols=5)
    tt.style = "Table Grid"
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["No.", "\uc774\uc804 \uc77c\uc790", "\uae08 \uc561 (\uc6d0)", "\uc131 \uaca9", "\uc785\uc99d \uc790\ub8cc"]
    for i, h in enumerate(headers):
        sc(tt.rows[0].cells[i], h, bold=True, sz=9)
    for r in range(1, rows+1):
        sc(tt.rows[r].cells[0], str(r), sz=9)
        sc(tt.rows[r].cells[1], "20    .    .    ", sz=9)
        sc(tt.rows[r].cells[2], "", sz=9)
        sc(tt.rows[r].cells[3], "", sz=9)
        sc(tt.rows[r].cells[4], "\uacc4\uc88c\uc774\uccb4 \ud655\uc778\uc11c", sz=9)
    sc(tt.rows[rows+1].cells[1], "\ud569   \uacc4", bold=True, sz=10)
    sc(tt.rows[rows+1].cells[2], "\uae08                    \uc6d0\uc815", bold=True, sz=10)

# ======================== CONTRACT ========================
def make_contract():
    d = init_doc()

    # Title
    t = d.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("\ubd80 \ub3d9 \uc0b0  \ub9e4 \ub9e4 \uacc4 \uc57d \uc11c")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(26, 35, 126)

    st = d.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("(\uac00\uc871\uac04 \uc800\uac00\ub9e4\ub9e4 - \uc0c1\uc99d\uc138\ubc95 \uc81c35\uc870 \uace0\ub824)")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    d.add_paragraph("")
    heading(d, "\u3010 \ub2f9 \uc0ac \uc790 \u3011")

    party_table(d, "\ub9e4\ub3c4\uc778(\uac11)")
    d.add_paragraph("")
    party_table(d, "\ub9e4\uc218\uc778(\uc744)")
    d.add_paragraph("")

    p = d.add_paragraph("\uac11\uacfc \uc744\uc740 \uc544\ub798 \ubd80\ub3d9\uc0b0\uc5d0 \uad00\ud558\uc5ec \ub2e4\uc74c\uacfc \uac19\uc774 \ub9e4\ub9e4\uacc4\uc57d\uc744 \uccb4\uacb0\ud55c\ub2e4.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("")

    # Property
    heading(d, "\u3010 \ubd80\ub3d9\uc0b0\uc758 \ud45c\uc2dc \u3011")
    tp = d.add_table(rows=5, cols=2)
    tp.style = "Table Grid"
    tp.alignment = WD_TABLE_ALIGNMENT.CENTER
    props = [
        ("\uc18c \uc7ac \uc9c0", "\uacbd\uae30\ub3c4 \uace0\uc591\uc2dc \uc77c\uc0b0\uc11c\uad6c                                    "),
        ("\ub2e8 \uc9c0 \uba85", "\ubc31\uc1a1\ud55c\uc2e0\uc544\ud30c\ud2b8          \ub3d9          \ud638"),
        ("\uba74    \uc801", "\uacf5\uae09\uba74\uc801 50B\u33a1  (\uc804\uc6a9\uba74\uc801 37B\u33a1)"),
        ("\uad6c    \uc870", "\ucca0\uadfc\ucf58\ud06c\ub9ac\ud2b8 / \uc544\ud30c\ud2b8"),
        ("\uc900\uacf5\ub144\ub3c4", "1992\ub144 9\uc6d4 (34\ub144\ucc28)"),
    ]
    for i, (a, b) in enumerate(props):
        sc(tp.rows[i].cells[0], a, bold=True)
        sc(tp.rows[i].cells[1], b)

    d.add_paragraph("")

    # Price table
    heading(d, "\u3010 \ub9e4\ub9e4\ub300\uae08 \ubc0f \uc9c0\uae09\ubc29\ubc95 \u3011")
    tpr = d.add_table(rows=5, cols=3)
    tpr.style = "Table Grid"
    tpr.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["\uad6c \ubd84", "\uae08 \uc561", "\uc9c0\uae09\uc77c / \ube44\uace0"]):
        sc(tpr.rows[0].cells[i], h, bold=True)
    price_rows = [
        ("\ub9e4\ub9e4\ub300\uae08 \ucd1d\uc561", "\uae08 \uc77c\uc5b5\uc624\ucc9c\ub9cc\uc6d0\uc815\n(150,000,000\uc6d0)", ""),
        ("\uacc4 \uc57d \uae08", "\uae08                    \uc6d0\uc815", "\uacc4\uc57d \uc2dc \uc9c0\uae09"),
        ("\uc0c1 \uacc4 \uae08", "\uae08                    \uc6d0\uc815", "\ubcc4\uccca \uc0c1\uacc4\ud569\uc758\uc11c\uc5d0 \uc758\ud568"),
        ("\uc794    \uae08", "\uae08                    \uc6d0\uc815", "20    \ub144    \uc6d4    \uc77c"),
    ]
    for i, (a, b, c) in enumerate(price_rows):
        sc(tpr.rows[i + 1].cells[0], a, bold=True)
        sc(tpr.rows[i + 1].cells[1], b)
        sc(tpr.rows[i + 1].cells[2], c)

    d.add_paragraph("")
    heading(d, "\u3010 \uacc4 \uc57d \uc870 \ud56d \u3011")

    # Clauses
    clauses = [
        ("\uc81c1\uc870 (\ub9e4\ub9e4\uac00\uc561 \uc0b0\uc815 \uadfc\uac70)", [
            "1. \ubcf8 \ub9e4\ub9e4\uac00\uc561\uc740 \uc0c1\uc18d\uc138 \ubc0f \uc99d\uc5ec\uc138\ubc95 \uc81c35\uc870(\uc800\uac00 \uc591\uc218\uc5d0 \ub530\ub978 \uc774\uc775\uc758 \uc99d\uc5ec)\uc758 \uae30\uc900\uc744 \uace0\ub824\ud558\uc5ec \ub2f9\uc0ac\uc790 \uac04 \ud569\uc758\ub85c \uc124\uc815\ud558\uc600\ub2e4.",
            "2. \uc2dc\uac00 \uc0b0\uc815 \uadfc\uac70: \ub3d9\uc77c \ub2e8\uc9c0 \uc720\uc0ac \uba74\uc801(50B\u33a1) \ucd5c\uadfc \uc2e4\uac70\ub798\uac00\ub97c \ucc38\uace0\ud558\uc600\uc73c\uba70, \ud574\ub2f9 \ud638\uc218\uc758 \uac1c\ubcc4 \uc0ac\uc815(\uce35\uc218, \ud5a5, \ub178\ud6c4\ub3c4, \ub0b4\ubd80 \uc0c1\ud0dc \ub4f1)\uc744 \ubc18\uc601\ud558\uc600\ub2e4.",
            "   - \ucd5c\uadfc \uc2e4\uac70\ub798\uac00: 2026.01 12\uce35 2\uc5b5 3,000\ub9cc\uc6d0 (\uad6d\ud1a0\uad50\ud1b5\ubd80 \uc2e4\uac70\ub798\uac00 \uacf5\uac1c\uc2dc\uc2a4\ud15c)",
            "   - \ucd5c\uadfc 3\uac1c\uc6d4 \ud3c9\uade0: \uc57d 2\uc5b5 4,600\ub9cc\uc6d0",
            "3. \uac10\uc815\ud3c9\uac00\uc11c\uac00 \uc788\ub294 \uacbd\uc6b0, \uac10\uc815\ud3c9\uac00\uc561\uc744 \uc2dc\uac00\ub85c \uc6b0\uc120 \uc801\uc6a9\ud55c\ub2e4. (\ubcc4\uccca)",
        ]),
        ("\uc81c2\uc870 (\ub300\uae08 \uc9c0\uae09 \ubc29\ubc95)", [
            "1. \uacc4\uc57d\uae08\uc740 \uacc4\uc57d \uccb4\uacb0 \uc2dc \ub9e4\uc218\uc778\uc774 \ub9e4\ub3c4\uc778\uc758 \uc9c0\uc815 \uacc4\uc88c\ub85c \uc774\uccb4\ud55c\ub2e4.",
            "2. \uc0c1\uacc4\uae08\uc740 \ub9e4\ub3c4\uc778\uc758 \ub9e4\uc218\uc778\uc5d0 \ub300\ud55c \uae30\uc874 \ucc44\ubb34(\ubcc4\uccca \ucc44\uad8c\ucc44\ubb34 \ud655\uc778\uc11c \ubc0f \uc0c1\uacc4 \ud569\uc758\uc11c \ucc38\uc870)\uc640 \uc0c1\uacc4 \ucc98\ub9ac\ud55c\ub2e4.",
            "3. \uc794\uae08\uc740 \uc704 \uc9c0\uae09\uc77c\uae4c\uc9c0 \ub9e4\ub3c4\uc778\uc758 \uc9c0\uc815 \uacc4\uc88c\ub85c \uc774\uccb4\ud558\uba70, \uc774\uccb4 \ud655\uc778\uc99d\uc744 \ubcf4\uad00\ud55c\ub2e4.",
            "4. \ubaa8\ub4e0 \ub300\uae08 \uc9c0\uae09\uc740 \ubc18\ub4dc\uc2dc \uacc4\uc88c\uc774\uccb4\ub85c \ud558\uba70, \ud604\uae08 \uc9c0\uae09\uc740 \ubd88\uc778\uc815\ud55c\ub2e4.",
        ]),
        ("\uc81c3\uc870 (\uc18c\uc720\uad8c \uc774\uc804)", [
            "1. \ub9e4\ub3c4\uc778\uc740 \uc794\uae08 \uc218\ub839\uacfc \ub3d9\uc2dc\uc5d0 \ub9e4\uc218\uc778 \uba85\uc758\ub85c \uc18c\uc720\uad8c\uc774\uc804\ub4f1\uae30\uc5d0 \ud544\uc694\ud55c \uc77c\uccb4\uc758 \uc11c\ub958\ub97c \uad50\ubd80\ud55c\ub2e4.",
            "2. \ub4f1\uae30\ube44\uc6a9(\ucde8\ub4dd\uc138, \ub4f1\ub85d\uba74\ud5c8\uc138, \ubc95\ubb34\uc0ac \uc218\uc218\ub8cc \ub4f1)\uc740 \ub9e4\uc218\uc778\uc774 \ubd80\ub2f4\ud55c\ub2e4.",
        ]),
        ("\uc81c4\uc870 (\uc784\ub300\ucc28 \uc2b9\uacc4)", [
            "1. \ubcf8 \ubd80\ub3d9\uc0b0\uc5d0 \uc124\uc815\ub41c \uc784\ub300\ucc28 \uad00\uacc4\ub294 \ub9e4\uc218\uc778\uc774 \uadf8\ub300\ub85c \uc2b9\uacc4\ud55c\ub2e4.",
            "   - \uc804\uc138\ubcf4\uc99d\uae08: \uae08                    \uc6d0",
            "   - \uc784\ub300\ucc28 \ub9cc\uae30\uc77c: 20    \ub144    \uc6d4    \uc77c",
            "2. \ub9e4\ub3c4\uc778\uc740 \uc784\ucc28\uc778\uc5d0\uac8c \uc784\ub300\uc778 \ubcc0\uacbd \uc0ac\uc2e4\uc744 \uc11c\uba74\uc73c\ub85c \ud1b5\uc9c0\ud55c\ub2e4.",
        ]),
        ("\uc81c5\uc870 (\ub2f4\ubcf4\uad8c \ubc0f \uc81c\ud55c\ubb3c\uad8c)", [
            "1. \ub9e4\ub3c4\uc778\uc740 \uc794\uae08 \uc218\ub839\uc77c\uae4c\uc9c0 \ubcf8 \ubd80\ub3d9\uc0b0\uc5d0 \uc124\uc815\ub41c \uadfc\uc800\ub2f9\uad8c, \uac00\uc555\ub958, \uac00\ucc98\ubd84 \ub4f1 \uc77c\uccb4\uc758 \ubd80\ub2f4\uc744 \ub9d0\uc18c\ud55c\ub2e4.",
        ]),
        ("\uc81c6\uc870 (\ud2b9\uc57d\uc0ac\ud56d)", [
            "1. \ub9e4\ub3c4\uc778\uacfc \ub9e4\uc218\uc778\uc740 \uc9c1\uacc4 \uac00\uc871(\ubd80\ub140) \uad00\uacc4\uc784\uc744 \ud655\uc778\ud55c\ub2e4.",
            "2. \ubcf8 \uac70\ub798\ub294 \uac00\uc871 \uac04 \uc720\uc0c1\uac70\ub798\uc774\uba70, \ub9e4\ub9e4\ub300\uae08 \uc804\uc561\uc774 \uc2e4\uc81c \uc9c0\uae09(\ud604\uae08 \uc774\uccb4 \ub610\ub294 \ucc44\ubb34 \uc0c1\uacc4)\ub428\uc744 \uc30d\ubc29 \ud655\uc778\ud55c\ub2e4.",
            "3. \uacfc\uac70 \uc790\uae08 \uad00\uacc4 \uc815\ub9ac\ub97c \uc704\ud55c \uc0c1\uacc4 \ub0b4\uc5ed\uc740 \ubcc4\uccca \ucc44\uad8c\ucc44\ubb34 \ud655\uc778\uc11c \ubc0f \uc0c1\uacc4 \ud569\uc758\uc11c\uc5d0 \ub530\ub978\ub2e4.",
            "4. \ubcf8 \uacc4\uc57d\uc5d0 \ub530\ub978 \uc99d\uc5ec\uc138 \uc2e0\uace0 \uc758\ubb34\uac00 \ubc1c\uc0dd\ud558\ub294 \uacbd\uc6b0, \ub9e4\uc218\uc778\uc774 \ubc95\uc815 \uae30\ud55c \ub0b4 \uc790\uc9c4 \uc2e0\uace0\ud55c\ub2e4.",
            "5. \uae30\ud0c0 \ubcf8 \uacc4\uc57d\uc5d0 \uba85\uc2dc\ub418\uc9c0 \uc54a\uc740 \uc0ac\ud56d\uc740 \ubbfc\ubc95 \ubc0f \uad00\ub828 \ubc95\ub839\uc5d0 \ub530\ub978\ub2e4.",
        ]),
    ]

    for title, itms in clauses:
        p = d.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        add_items(d, itms)

    d.add_paragraph("")
    heading(d, "\u3010 \ubcc4 \uccca \uc11c \ub958 \u3011")
    add_items(d, [
        "1. \ucc44\uad8c\ucc44\ubb34 \ud655\uc778\uc11c \ubc0f \uc0c1\uacc4 \ud569\uc758\uc11c",
        "2. \uacc4\uc88c \uc774\uccb4 \ub0b4\uc5ed (\uacfc\uac70 \uc790\uae08 \ud750\ub984 \uc99d\ube59)",
        "3. \uac10\uc815\ud3c9\uac00\uc11c (\ud574\ub2f9 \uc2dc)",
        "4. \ubd80\ub3d9\uc0b0 \ub4f1\uae30\ubd80\ub4f1\ubcf8",
        "5. \uac74\ucd95\ubb3c\ub300\uc7a5",
        "6. \uc784\ub300\ucc28\uacc4\uc57d\uc11c \uc0ac\ubcf8 (\ud574\ub2f9 \uc2dc)",
    ])

    d.add_paragraph("")
    d.add_paragraph("")
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\uc704 \uacc4\uc57d \ub0b4\uc6a9\uc744 \ud655\uc778\ud558\uace0 \uc774\uc5d0 \uc11c\uba85 \ub0a0\uc778\ud55c\ub2e4.")
    r.font.size = Pt(11)
    d.add_paragraph("")

    sign_block(d, ["\ub9e4\ub3c4\uc778 (\uac11)", "\ub9e4\uc218\uc778 (\uc744)"])

    path = os.path.join(OUT, "\ub9e4\ub9e4\uacc4\uc57d\uc11c_\ubc31\uc1a1\ud55c\uc2e0_\uc800\uac00\ub9e4\ub9e4.docx")
    d.save(path)
    sz = os.path.getsize(path)
    print(f"Contract saved: {sz:,} bytes")
    return path


# ======================== AGREEMENT ========================
def make_agreement():
    d = init_doc()

    t = d.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("\ucc44\uad8c\ucc44\ubb34 \ud655\uc778\uc11c \ubc0f \uc0c1\uacc4 \ud569\uc758\uc11c")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(26, 35, 126)

    d.add_paragraph("")
    heading(d, "\u3010 \ub2f9 \uc0ac \uc790 \u3011")

    tb = d.add_table(rows=4, cols=4)
    tb.style = "Table Grid"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    party_data = [
        ("\ucc44\ubb34\uc790 (\uac11)", "\uc131 \uba85", "", "\uc8fc\ubbfc\ubc88\ud638:       -       "),
        ("", "\uc8fc \uc18c", "", ""),
        ("\ucc44\uad8c\uc790 (\uc744)", "\uc131 \uba85", "", "\uc8fc\ubbfc\ubc88\ud638:       -       "),
        ("", "\uc8fc \uc18c", "", ""),
    ]
    for i, (a, b, c, dd) in enumerate(party_data):
        sc(tb.rows[i].cells[0], a, bold=bool(a))
        sc(tb.rows[i].cells[1], b)
        sc(tb.rows[i].cells[2], c)
        sc(tb.rows[i].cells[3], dd)

    d.add_paragraph("")
    d.add_paragraph("\uac11(\ucc44\ubb34\uc790, \uc544\ubc84\uc9c0)\uacfc \uc744(\ucc44\uad8c\uc790, \ub538)\uc740 \uacfc\uac70 \uc790\uae08 \uac70\ub798 \uad00\uacc4\ub97c \uc544\ub798\uc640 \uac19\uc774 \ud655\uc778\ud558\uace0, \ubc31\uc1a1\ud55c\uc2e0\uc544\ud30c\ud2b8 \ub9e4\ub9e4\ub300\uae08\uacfc\uc758 \uc0c1\uacc4 \ucc98\ub9ac\uc5d0 \ud569\uc758\ud55c\ub2e4.")
    d.add_paragraph("")

    # Art 1
    heading(d, "\uc81c1\uc870 (\ucc44\uad8c\ucc44\ubb34\uc758 \ubc1c\uc0dd \uacbd\uc704)", 12)
    add_items(d, [
        "1. \uc744(\ub538)\uc740 \uacfc\uac70 \uc218\ub144\uac04 \uac11(\uc544\ubc84\uc9c0)\uc5d0\uac8c \ud22c\uc790 \uc694\uccad, \ub300\uc5ec, \uc0dd\ud65c\ube44 \uc9c0\uc6d0 \ub4f1 \ub2e4\uc591\ud55c \uba85\ubaa9\uc73c\ub85c \uc790\uae08\uc744 \uc774\uc804\ud558\uc600\ub2e4.",
        "2. \uac11(\uc544\ubc84\uc9c0) \uc5ed\uc2dc \uc744(\ub538)\uc5d0\uac8c \uc77c\ubd80 \uc790\uae08\uc744 \uc774\uc804\ud55c \uc0ac\uc2e4\uc774 \uc788\ub2e4.",
        "3. \uc591 \ub2f9\uc0ac\uc790\ub294 \uc0c1\uae30 \uc790\uae08 \uac70\ub798\uc758 \uc131\uaca9(\ud22c\uc790, \ub300\uc5ec, \uc99d\uc5ec \ub4f1)\uc774 \ud63c\uc7ac\ub418\uc5b4 \uc788\uc74c\uc744 \uc778\uc815\ud558\uba70, \ubcf8 \ud569\uc758\uc11c\ub97c \ud1b5\ud574 \uc774\ub97c \uc815\ub9ac\ud55c\ub2e4.",
    ])
    d.add_paragraph("")

    # Art 2 & 3 - Transfer tables
    heading(d, "\uc81c2\uc870 (\uc744 \u2192 \uac11 \uc790\uae08 \uc774\uc804 \ub0b4\uc5ed)", 12)
    d.add_paragraph("\uc744(\ub538)\uc774 \uac11(\uc544\ubc84\uc9c0)\uc5d0\uac8c \uc774\uc804\ud55c \uae08\uc561:")
    make_transfer_table(d)
    d.add_paragraph("")

    heading(d, "\uc81c3\uc870 (\uac11 \u2192 \uc744 \uc790\uae08 \uc774\uc804 \ub0b4\uc5ed)", 12)
    d.add_paragraph("\uac11(\uc544\ubc84\uc9c0)\uc774 \uc744(\ub538)\uc5d0\uac8c \uc774\uc804\ud55c \uae08\uc561:")
    make_transfer_table(d)
    d.add_paragraph("")

    # Art 4 - Net amount
    heading(d, "\uc81c4\uc870 (\uc21c\ucc44\uad8c\uc561 \ud655\uc815)", 12)
    t4 = d.add_table(rows=4, cols=2)
    t4.style = "Table Grid"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    net_data = [
        ("\uc744\u2192\uac11 \uc774\uc804 \ucd1d\uc561 (A)", "\uae08                              \uc6d0\uc815"),
        ("\uac11\u2192\uc744 \uc774\uc804 \ucd1d\uc561 (B)", "\uae08                              \uc6d0\uc815"),
        ("\uc21c\ucc44\uad8c\uc561 (A - B)", "\uae08                              \uc6d0\uc815"),
        ("\uc0c1\uacc4 \uc801\uc6a9 \uae08\uc561", "\uae08                              \uc6d0\uc815"),
    ]
    for i, (a, b) in enumerate(net_data):
        sc(t4.rows[i].cells[0], a, bold=True)
        sc(t4.rows[i].cells[1], b)
    d.add_paragraph("")

    # Art 5 - Crypto
    heading(d, "\uc81c5\uc870 (\ube44\ud2b8\ucf54\uc778/\uac70\ub798\uc18c \uacf5\ub3d9 \uc790\uae08 \uc6b4\uc6a9 \uc815\ub9ac)", 12)
    add_items(d, [
        "1. \uac11\uacfc \uc744\uc740 \uc544\ub798 \uac70\ub798\uc18c\uc5d0\uc11c \uacf5\ub3d9\uc73c\ub85c \uc790\uae08\uc744 \uc6b4\uc6a9\ud55c \uc0ac\uc2e4\uc744 \ud655\uc778\ud55c\ub2e4.",
        "   - \uac70\ub798\uc18c\uba85:                              (\uc608: \ubc14\uc774\ub09c\uc2a4, \uc5c5\ube44\ud2b8 \ub4f1)",
        "   - \uc6b4\uc6a9 \uae30\uac04: 20    \ub144    \uc6d4 ~ 20    \ub144    \uc6d4",
    ])
    d.add_paragraph("2. \uc6d0\uae08 \ud22c\uc785 \ub0b4\uc5ed:")
    t5 = d.add_table(rows=3, cols=3)
    t5.style = "Table Grid"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc(t5.rows[0].cells[0], "\uad6c\ubd84", bold=True)
    sc(t5.rows[0].cells[1], "\ud22c\uc785\uc561 (\uc6d0)", bold=True)
    sc(t5.rows[0].cells[2], "\uc785\uc99d \uc790\ub8cc", bold=True)
    sc(t5.rows[1].cells[0], "\uac11 (\uc544\ubc84\uc9c0)")
    sc(t5.rows[1].cells[1], "\uae08                    \uc6d0")
    sc(t5.rows[1].cells[2], "\uac70\ub798\uc18c \uc785\uae08 \ub0b4\uc5ed")
    sc(t5.rows[2].cells[0], "\uc744 (\ub538)")
    sc(t5.rows[2].cells[1], "\uae08                    \uc6d0")
    sc(t5.rows[2].cells[2], "\uac70\ub798\uc18c \uc785\uae08 \ub0b4\uc5ed")
    d.add_paragraph("")
    d.add_paragraph("3. \uc6b4\uc6a9 \uacb0\uacfc:")
    t6 = d.add_table(rows=5, cols=2)
    t6.style = "Table Grid"
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    result_data = [
        ("\ud604\uc7ac \ud3c9\uac00\uc561", "\uae08                    \uc6d0 (\uae30\uc900\uc77c: 20    \ub144    \uc6d4    \uc77c)"),
        ("\ucd1d \uc190\uc775", "[ ] \uc218\uc775 / [ ] \uc190\uc2e4   \uae08                    \uc6d0"),
        ("\uc6d4\ubcc4 \ucd5c\uace0 \uc218\uc775", "        \ub144     \uc6d4 (\uae08                    \uc6d0)"),
        ("\uc6d4\ubcc4 \ucd5c\ub300 \uc190\uc2e4", "        \ub144     \uc6d4 (\uae08                    \uc6d0)"),
        ("\uc815\uc0b0 \ubc29\ubc95", "[ ] \uc9c0\ubd84 \ube44\uc728 \uc815\uc0b0 / [ ] \ud569\uc758\uc11c \ud3ec\ud568 / [ ] \uae30\ud0c0"),
    ]
    for i, (a, b) in enumerate(result_data):
        sc(t6.rows[i].cells[0], a, bold=True)
        sc(t6.rows[i].cells[1], b)
    d.add_paragraph("")

    # Art 6 - Offset
    heading(d, "\uc81c6\uc870 (\uc0c1\uacc4 \ud569\uc758)", 12)
    add_items(d, [
        "1. \uac11\uacfc \uc744\uc740 \uc81c4\uc870\uc758 \uc21c\ucc44\uad8c\uc561 \uc911 \uae08                    \uc6d0\uc744 \ubc31\uc1a1\ud55c\uc2e0\uc544\ud30c\ud2b8 \ub9e4\ub9e4\ub300\uae08(\uae08 \uc77c\uc5b5\uc624\ucc9c\ub9cc\uc6d0\uc815)\uacfc \uc0c1\uacc4 \ucc98\ub9ac\ud558\uae30\ub85c \ud569\uc758\ud55c\ub2e4.",
        "2. \uc0c1\uacc4 \ud6c4 \ub9e4\uc218\uc778(\uc744)\uc774 \ucd94\uac00\ub85c \uc9c0\uae09\ud560 \uc794\uc561\uc740 \uae08                    \uc6d0\uc774\uba70, \uc774\ub294 \ub9e4\ub9e4\uacc4\uc57d\uc11c \uc81c2\uc870\uc5d0 \ub530\ub77c \uc9c0\uae09\ud55c\ub2e4.",
        "3. \ubcf8 \uc0c1\uacc4 \ucc98\ub9ac\ub85c \uc81c4\uc870\uc758 \uc21c\ucc44\uad8c\uc561 \uc911 \uc0c1\uacc4 \uc801\uc6a9 \uae08\uc561\uc5d0 \ud574\ub2f9\ud558\ub294 \ucc44\uad8c\uc740 \uc18c\uba78\ud55c\ub2e4.",
    ])
    d.add_paragraph("")

    # Art 7
    heading(d, "\uc81c7\uc870 (\ud655\uc778 \ubc0f \ubcf4\uc99d)", 12)
    add_items(d, [
        "1. \uac11\uacfc \uc744\uc740 \ubcf8 \ud655\uc778\uc11c\uc5d0 \uae30\uc7ac\ub41c \uae08\uc561 \ubc0f \uac70\ub798 \ub0b4\uc5ed\uc774 \uc0ac\uc2e4\uacfc \ubd80\ud569\ud568\uc744 \ud655\uc778\ud55c\ub2e4.",
        "2. \ubcf8 \ud655\uc778\uc11c\ub294 \ud5a5\ud6c4 \uc138\ubb34 \uc870\uc0ac \ub610\ub294 \ubc95\uc801 \ubd84\uc7c1 \uc2dc \uc99d\uac70 \uc790\ub8cc\ub85c \uc0ac\uc6a9\ub420 \uc218 \uc788\ub2e4.",
        "3. \ubcf8 \ud655\uc778\uc11c\uc5d0 \uae30\uc7ac\ub418\uc9c0 \uc54a\uc740 \ucd94\uac00 \ucc44\uad8c\ucc44\ubb34 \uad00\uacc4\ub294 \uc5c6\uc74c\uc744 \uc0c1\ud638 \ud655\uc778\ud55c\ub2e4.",
        "4. \ud5c8\uc704 \uae30\uc7ac\ub85c \uc778\ud55c \ubc95\uc801 \ucc45\uc784\uc740 \ud574\ub2f9 \ub2f9\uc0ac\uc790\uac00 \ubd80\ub2f4\ud55c\ub2e4.",
    ])
    d.add_paragraph("")

    heading(d, "\u3010 \ubcc4 \uccca \uc99d \ube59 \uc11c \ub958 \u3011")
    add_items(d, [
        "1. \uc740\ud589 \uacc4\uc88c \uac70\ub798\ub0b4\uc5ed \ud655\uc778\uc11c (\uac01 \uc740\ud589\ubcc4)",
        "2. \uac70\ub798\uc18c(\ubc14\uc774\ub09c\uc2a4 \ub4f1) \uc785\ucd9c\uae08 \ub0b4\uc5ed\uc11c",
        "3. \uac70\ub798\uc18c \uc190\uc775 \ub9ac\ud3ec\ud2b8 / \uc794\uace0 \uc2a4\ud06c\ub9b0\uc0f7",
        "4. \uae30\uc874 \ucc28\uc6a9\uc99d \ub610\ub294 \uc57d\uc815\uc11c (\uc788\ub294 \uacbd\uc6b0)",
        "5. \uae30\ud0c0 \uc790\uae08 \uc774\ub3d9 \uc99d\ube59 \uc790\ub8cc",
    ])

    d.add_paragraph("")
    d.add_paragraph("")
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\uc704 \ub0b4\uc6a9\uc774 \uc0ac\uc2e4\uacfc \ub2e4\ub984\uc5c6\uc74c\uc744 \ud655\uc778\ud558\uace0, \uc0c1\uacc4 \ucc98\ub9ac\uc5d0 \ud569\uc758\ud558\uc5ec \uc11c\uba85 \ub0a0\uc778\ud55c\ub2e4.")
    r.font.size = Pt(11)
    d.add_paragraph("")

    sign_block(d, ["\ucc44\ubb34\uc790 (\uac11)", "\ucc44\uad8c\uc790 (\uc744)"])

    d.add_paragraph("")
    heading(d, "\u3010 \uacf5 \uc99d \ub780 \u3011 (\uc120\ud0dd\uc0ac\ud56d)", 12)
    tn = d.add_table(rows=4, cols=2)
    tn.style = "Table Grid"
    tn.alignment = WD_TABLE_ALIGNMENT.CENTER
    notary = [
        ("\uacf5\uc99d\uc778", ""),
        ("\uacf5\uc99d\uc77c", "20      \ub144      \uc6d4      \uc77c"),
        ("\uacf5\uc99d\ubc88\ud638", ""),
        ("\uacf5\uc99d\uc778 \uc11c\uba85", "                                        (\uc9c1\uc778)"),
    ]
    for i, (a, b) in enumerate(notary):
        sc(tn.rows[i].cells[0], a, bold=True)
        sc(tn.rows[i].cells[1], b)

    path = os.path.join(OUT, "\uc0c1\uacc4\ud569\uc758\uc11c_\ucc44\uad8c\ucc44\ubb34\ud655\uc778\uc11c.docx")
    d.save(path)
    sz = os.path.getsize(path)
    print(f"Agreement saved: {sz:,} bytes")
    return path


if __name__ == "__main__":
    p1 = make_contract()
    p2 = make_agreement()
    print(f"\nDone! Files at:\n  {p1}\n  {p2}")
